import streamlit as st
import google.generativeai as genai
from PIL import Image
import tempfile
import os
import io
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- 1. CẤU HÌNH TRANG ---
st.set_page_config(page_title="Trợ lý Soạn Giáo Án 4.0", page_icon="📘", layout="centered")

FILE_KHUNG_NANG_LUC = "khungnanglucso.pdf"

# --- 2. HÀM XỬ LÝ WORD (GIỮ NGUYÊN NHƯ CŨ) ---
def add_formatted_text(paragraph, text):
    paragraph.style.font.name = 'Times New Roman'
    paragraph.style.font.size = Pt(14)
    parts = re.split(r'(\*\*.*?\*\*)', text) 
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            clean = part[2:-2]; run = paragraph.add_run(clean); run.bold = True
        else: run = paragraph.add_run(part)
        run.font.name = 'Times New Roman'; run.font.size = Pt(14)

def create_doc_stable(content, ten_bai, lop):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21); section.page_height = Cm(29.7)
    section.top_margin = Cm(2); section.bottom_margin = Cm(2); section.left_margin = Cm(3); section.right_margin = Cm(1.5)
    style = doc.styles['Normal']; font = style.font; font.name = 'Times New Roman'; font.size = Pt(14); style.paragraph_format.line_spacing = 1.2
    
    head = doc.add_heading(f'KẾ HOẠCH BÀI DẠY: {ten_bai.upper()}', 0); head.alignment = 1 
    for run in head.runs: run.font.name = 'Times New Roman'; run.font.size = Pt(16); run.font.bold = True; run.font.color.rgb = RGBColor(0, 0, 0)
    p_lop = doc.add_paragraph(f'Lớp: {lop}'); p_lop.alignment = 1; p_lop.runs[0].bold = True; p_lop.runs[0].font.name = 'Times New Roman'; p_lop.runs[0].font.size = Pt(14)
    doc.add_paragraph("-" * 60).alignment = 1
    
    lines = content.split('\n'); i = 0
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith('#'): line = line.replace('#', '').strip()
        if line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'): table_lines.append(lines[i].strip()); i += 1
            if len(table_lines) >= 3: 
                try:
                    valid_rows = [r for r in table_lines if '---' not in r]
                    if valid_rows:
                        cols_count = len(valid_rows[0].split('|')) - 2 
                        if cols_count > 0:
                            table = doc.add_table(rows=len(valid_rows), cols=cols_count); table.style = 'Table Grid'; table.autofit = True
                            for r_idx, r_text in enumerate(valid_rows):
                                cells_data = r_text.split('|')[1:-1]
                                for c_idx, cell_text in enumerate(cells_data):
                                    if c_idx < cols_count:
                                        cell = table.cell(r_idx, c_idx); cell._element.clear_content()
                                        raw_content = cell_text.strip().replace('<br>', '\n').replace('<br/>', '\n')
                                        sub_lines = raw_content.split('\n')
                                        for sub_line in sub_lines:
                                            sub_line = sub_line.strip(); 
                                            if not sub_line: continue
                                            p = cell.add_paragraph(); p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(2); p.paragraph_format.line_spacing = 1.1
                                            if r_idx == 0: p.alignment = 1; run = p.add_run(sub_line.replace('**','')); run.bold = True; run.font.name = 'Times New Roman'; run.font.size = Pt(14)
                                            else: add_formatted_text(p, sub_line)
                except: pass
            continue 
        if not line: i += 1; continue
        if re.match(r'^(I\.|II\.|III\.|IV\.|V\.)', line) or (re.match(r'^\d+\.', line) and len(line) < 50): 
            clean = line.replace('**', '').strip(); p = doc.add_paragraph(clean); p.runs[0].bold = True; p.runs[0].font.name = 'Times New Roman'; p.runs[0].font.size = Pt(14)
        elif line.startswith('- ') or line.startswith('* '):
            clean = line[2:].strip(); p = doc.add_paragraph(); run_dash = p.add_run("- "); run_dash.font.name = 'Times New Roman'; run_dash.font.size = Pt(14); add_formatted_text(p, clean)
        else: p = doc.add_paragraph(); add_formatted_text(p, line)
        i += 1
    return doc

# --- 3. GIAO DIỆN CHÍNH ---
st.markdown("""
<div style='text-align: center; background: linear-gradient(135deg, #004e92 0%, #000428 100%); padding: 20px; border-radius: 10px; color: white; margin-bottom: 20px;'>
    <h1 style='color: white; margin:0;'>📘 TRỢ LÝ SOẠN GIÁO ÁN NLS</h1>
    <p style='margin-top:10px;'>Tác giả: Thầy Lý Văn Du - PTDTBT TH Du Tiến </p>
</div>
""", unsafe_allow_html=True)

# --- [PHẦN QUAN TRỌNG NHẤT: BẮT BUỘC NHẬP KEY] ---
with st.sidebar:
    st.header("🔐 Đăng nhập hệ thống")
    st.info("Để sử dụng, thầy/cô vui lòng nhập API Key cá nhân của mình.")
    
    api_key = st.text_input("🔑 Nhập API Key (Google AI Studio):", type="password")
    
    st.markdown("---")
    with st.expander("❓ Chưa có Key? Xem hướng dẫn"):
        st.markdown("""
        1. Truy cập [aistudio.google.com](https://aistudio.google.com/)
        2. Đăng nhập Gmail
        3. Chọn **Get API key** -> **Create API key**
        4. Copy mã và dán vào ô bên trên.
        """)

# [CHẶN]: Nếu không có Key -> Dừng lại luôn, không hiện nội dung bên dưới
if not api_key:
    st.warning("⬅️ Vui lòng nhập API Key vào thanh bên trái để bắt đầu soạn bài!")
    st.stop() # Lệnh này làm App dừng lại tại đây

# Nếu có Key -> Cấu hình và chạy tiếp
try:
    genai.configure(api_key=api_key)
except Exception as e:
    st.error(f"API Key không hợp lệ! Vui lòng kiểm tra lại. Lỗi: {e}")
    st.stop()

# --- NỘI DUNG CHÍNH (CHỈ HIỆN KHI ĐÃ CÓ KEY) ---

# 1. TÀI LIỆU
st.markdown('### 📂 1. TÀI LIỆU NGUỒN')
has_framework = False
if os.path.exists(FILE_KHUNG_NANG_LUC):
    st.success(f"✅ Đã tự động tích hợp: {FILE_KHUNG_NANG_LUC}")
    has_framework = True
else:
    st.info(f"ℹ️ Chưa có file '{FILE_KHUNG_NANG_LUC}' trong hệ thống.")

uploaded_files = st.file_uploader("Tải Ảnh/PDF bài dạy (Kéo thả vào đây):", type=["jpg", "png", "pdf"], accept_multiple_files=True)
if uploaded_files:
    cols = st.columns(3)
    for i, f in enumerate(uploaded_files):
        if f.type in ["image/jpeg", "image/png"]:
            with cols[i%3]: st.image(f, caption=f.name)
        else:
            with cols[i%3]: st.info(f"📄 {f.name}")

# 2. THÔNG TIN
st.markdown('### 📝 2. THÔNG TIN BÀI DẠY')
c1, c2 = st.columns(2)
with c1: lop = st.text_input("📚 Lớp:", "Lớp 4")
with c2: ten_bai = st.text_input("📌 Tên bài học:", placeholder="Ví dụ: Học hát bài...")
noidung_bosung = st.text_area("✍️ Ghi chú thêm:", height=100)
yeu_cau_them = st.text_input("💡 Yêu cầu đặc biệt:", placeholder="Ví dụ: Tích hợp trò chơi khởi động...")

# 3. NÚT XỬ LÝ
st.markdown("<br>", unsafe_allow_html=True)
if st.button("🚀 SOẠN GIÁO ÁN NGAY", type="primary", use_container_width=True):
    if not uploaded_files and not noidung_bosung and not has_framework:
        st.toast("Thiếu tài liệu! Hãy tải ảnh SGK lên.", icon="⚠️")
    else:
        try:
            with st.spinner('AI đang soạn giáo án (Model Gemini 1.5 Flash)...'):
                model = genai.GenerativeModel('gemini-2.5-flash-lite-preview-09-2025')
                
                # PROMPT GIỮ NGUYÊN THEO YÊU CẦU CỦA THẦY
                prompt_instruction = f"""
                Đóng vai là một Giáo viên Tiểu học giỏi, am hiểu chương trình GDPT 2018.
                Nhiệm vụ: Soạn Kế hoạch bài dạy (Giáo án) cho bài: "{ten_bai}" - {lop}.

                DỮ LIỆU ĐẦU VÀO:
                - Hãy phân tích TẤT CẢ các tài liệu hình ảnh và PDF được đính kèm để lấy nội dung kiến thức.
                - (Nếu có) File PDF Khung năng lực số đính kèm: Hãy dùng để đối chiếu nội dung bài học và đưa vào mục Năng lực số.
                - Kết hợp với ghi chú bổ sung: "{noidung_bosung}".

                YÊU CẦU LUÔN LUÔN TUÂN THỦ CẤU TRÚC (CÔNG VĂN 2345):
                I. Yêu cầu cần đạt: Trong phần này lại chia thành các phần sau: 
                1. Học sinh thực hiện được, 
                2. Học sinh vận dụng được, 
                3. Phát triển năng lực (bao gồm năng lực đặc thù, năng lực chung, phát triển năng lực số), 
                4. Phát triển phẩm chất.
                * Nội dung tích hợp (VD: Học thông qua chơi, Công dân số,...)
                    - Lưu ý: Thêm phát triển năng lực số trong mục phát triển năng lực (Dựa vào file Khung năng lực nếu có).
                    - Nội dung tích hợp Học thông qua chơi trong Yêu cầu cần đạt cần cụ thể chi tiết hơn chút nữa.
                    - Nội dung tích hợp Công dân số cũng cần cụ thể hơn trong yêu cầu cần đạt.

                II. Đồ dùng dạy học
                1. Giáo viên
                2. Học sinh

                III. Tiến trình dạy học
                [QUAN TRỌNG] PHẦN NÀY PHẢI TRÌNH BÀY DƯỚI DẠNG BẢNG (MARKDOWN TABLE) 2 CỘT. 
                TÊN CÁC HOẠT ĐỘNG PHẢI NẰM Ở CỘT 1 (HOẠT ĐỘNG GIÁO VIÊN).

                | HOẠT ĐỘNG CỦA GIÁO VIÊN | HOẠT ĐỘNG CỦA HỌC SINH |
                |---|---|
                | **1. Hoạt động 1 - Khởi động:**<br>- GV tổ chức... | - HS tham gia... |
                | **2. Hoạt động 2 - Hình thành kiến thức mới:**<br>- GV hướng dẫn... | - HS quan sát... |
                | **3. Hoạt động 3 - Thực hành - luyện tập:**<br>- GV yêu cầu... | - HS thực hiện... |
                | **4. Hoạt động 4 - Vận dụng:**<br>- GV gợi mở... | - HS chia sẻ... |

                YÊU CẦU CHI TIẾT CHO TIẾN TRÌNH DẠY HỌC:
                - Cần chi tiết cụ thể (đặc biệt là Hoạt động của học sinh).
                - Các ý trong tiến trình dạy học được bắt đầu bằng dấu gạch đầu dòng (-).
                - Tích hợp Học thông qua chơi vào 1 số hoạt động phù hợp.
                - Riêng các trò chơi trong tiến trình dạy học cần TRÌNH BÀY RÕ LUẬT CHƠI.
                - Không cần ghi "Mục tiêu HTQC".
                - Tiết học chỉ có 35 phút, hãy điều chỉnh lượng kiến thức và hoạt động hợp lý.
                - Không kèm chú thích nguồn trong bài soạn.
                - Tuyệt đối chỉ bao gồm 4 Hoạt động, không phát sinh thêm.
                - LUÔN LUÔN TUÂN THỦ THEO NHỮNG YÊU CẦU TRÊN

                IV. Điều chỉnh sau tiết dạy

                Lưu ý chung: Bám sát nội dung trong Sách giáo khoa và sách giáo viên (từ tài liệu đính kèm) để đưa nội dung vào bài soạn cho chính xác. KHÔNG dùng ký tự # ở đầu dòng.

                LƯU Ý QUAN TRỌNG TỪ NGƯỜI DÙNG: {yeu_cau_them}
                """

                input_data = [prompt_instruction]
                temp_paths = []
                
                if has_framework: input_data.append(genai.upload_file(FILE_KHUNG_NANG_LUC))
                
                if uploaded_files:
                    for f in uploaded_files:
                        if f.type == "application/pdf":
                            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                                tmp.write(f.getvalue())
                                temp_paths.append(tmp.name)
                            input_data.append(genai.upload_file(tmp.name))
                        else:
                            input_data.append(Image.open(f))
                
                if noidung_bosung: input_data.append(noidung_bosung)
                
                response = model.generate_content(input_data)
                
                st.markdown("### 📄 KẾT QUẢ BÀI SOẠN:")
                st.markdown(f'<div class="lesson-plan-paper">{response.text}</div>', unsafe_allow_html=True)
                
                doc = create_doc_stable(response.text, ten_bai, lop)
                buf = io.BytesIO()
                doc.save(buf)
                buf.seek(0)
                
                st.download_button(
                    label="⬇️ TẢI FILE WORD (.DOCX) CHUẨN A4, CĂN LỀ",
                    data=buf,
                    file_name=f"GiaoAn_{ten_bai}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary"
                )
                
                for p in temp_paths: os.remove(p)

        except Exception as e:
            st.error(f"Có lỗi xảy ra: {e}")

# --- CHÂN TRANG ---
st.markdown("---")
st.markdown("<div style='text-align: center; color: #666;'>© 2025 – 
Lý Văn Du - Trường PTDTBT Tiểu học Du Tiến - ĐT: 0964 973 866 </div>", unsafe_allow_html=True)

